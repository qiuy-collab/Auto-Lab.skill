import argparse
import json
from pathlib import Path

from docx import Document


def parse_args():
    parser = argparse.ArgumentParser(description="Initialize an auto-lab run directory.")
    parser.add_argument("--requirements", required=True, help="Path to the requirement document.")
    parser.add_argument("--template", required=True, help="Path to the docx template.")
    parser.add_argument("--output-dir", required=True, help="Directory for the generated run files.")
    parser.add_argument("--output-docx-name", required=True, help="Final output docx file name.")
    return parser.parse_args()


def analyze_template(template_path: Path):
    doc = Document(str(template_path))
    paragraphs = list(doc.paragraphs)
    non_empty = [p.text.strip() for p in paragraphs if p.text.strip()]
    table_shapes = []
    for table in doc.tables:
        rows = len(table.rows)
        cols = max((len(row.cells) for row in table.rows), default=0)
        table_shapes.append({"rows": rows, "cols": cols})
    return {
        "paragraph_count": len(paragraphs),
        "non_empty_paragraph_count": len(non_empty),
        "table_count": len(doc.tables),
        "table_shapes": table_shapes,
        "sample_headings_or_text": non_empty[:20],
    }


def write_script_stub(path: Path, script_role: str):
    stub = f'''import argparse
import json
from pathlib import Path

AUTO_LAB_TEMPLATE_SCRIPT_STUB = True


def parse_args():
    parser = argparse.ArgumentParser(description="Template-specific {script_role} script for the current docx task.")
'''
    if script_role == "fill":
        stub += """    parser.add_argument("--template", required=True, help="Path to the template docx")
    parser.add_argument("--copywriting", required=True, help="Path to copywriting.md")
    parser.add_argument("--output", required=True, help="Path to the output docx")
    parser.add_argument("--requirement-analysis", default=None, help="Path to requirement_analysis.json")
    parser.add_argument("--pre-task-plan", default=None, help="Path to pre_task_plan.json")
"""
    elif script_role == "insert":
        stub += """    parser.add_argument("--docx", required=True, help="Path to the docx to insert images into")
    parser.add_argument("--insert-config", required=True, help="Path to insert_config.json")
    parser.add_argument("--images-dir", default=None, help="Path to generated_images directory")
"""
    else:
        stub += """    parser.add_argument("--docx", required=True, help="Path to the docx to verify")
    parser.add_argument("--insert-config", required=True, help="Path to insert_config.json")
"""
    stub += """    return parser.parse_args()


def main():
    raise SystemExit(
        "This is a template-specific script stub. Replace it with logic for the current template.\\n"
        "Read the template structure and customize this script to:\\n"
        "  - fill: parse the template, fill in content from copywriting.md, respect template formatting\\n"
        "  - insert: read insert_config.json, place images at the correct positions\\n"
        "  - verify: check that all placeholders are filled and images are placed correctly\\n"
        "Use vendor/minimax-docx/SKILL.md for complex DOCX operations."
    )


if __name__ == "__main__":
    main()
"""
    path.write_text(stub, encoding="utf-8")


def main():
    args = parse_args()
    requirements = Path(args.requirements).expanduser().resolve()
    template = Path(args.template).expanduser().resolve()
    if template.suffix.lower() != ".docx":
        raise SystemExit(
            f"auto-lab currently requires a .docx template for python-docx processing. "
            f"Please convert this template first: {template}"
        )

    output_dir = Path(args.output_dir).expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    images_dir = output_dir / "generated_images"
    images_dir.mkdir(exist_ok=True)
    task_scripts_dir = output_dir / "task_scripts"
    task_scripts_dir.mkdir(exist_ok=True)

    output_docx = output_dir / args.output_docx_name
    workflow_path = output_dir / "workflow.json"
    checklist_path = output_dir / "requirement_checklist.json"
    requirement_analysis_path = output_dir / "requirement_analysis.json"
    pre_task_plan_path = output_dir / "pre_task_plan.json"
    copywriting_path = output_dir / "copywriting.md"
    prompt_config_path = output_dir / "prompt_config.json"
    browser_capture_plan_path = output_dir / "browser_capture_plan.json"
    diagram_plan_path = output_dir / "diagram_plan.json"
    video_plan_path = output_dir / "video_plan.json"
    reference_template_cleanup_path = output_dir / "reference_template_cleanup.json"
    submission_package_path = output_dir / "submission_package.json"
    insert_config_path = output_dir / "insert_config.json"
    template_manifest_path = output_dir / "template_manifest.json"
    delivery_review_path = output_dir / "delivery_review.json"
    fill_script_path = task_scripts_dir / "fill_template.py"
    insert_script_path = task_scripts_dir / "insert_images.py"
    verify_script_path = task_scripts_dir / "verify_template.py"

    template_manifest = analyze_template(template)
    template_manifest["template_path"] = str(template)
    template_manifest["template_protection_rule"] = (
        "Do not change original structure, fixed text, or template styling unless a position is explicitly identified as fillable."
    )

    workflow = {
        "requirements_path": str(requirements),
        "template_path": str(template),
        "output_dir": str(output_dir),
        "output_docx": str(output_docx),
        "requirement_checklist_path": str(checklist_path),
        "requirement_analysis_path": str(requirement_analysis_path),
        "pre_task_plan_path": str(pre_task_plan_path),
        "copywriting_path": str(copywriting_path),
        "prompt_config_path": str(prompt_config_path),
        "browser_capture_plan_path": str(browser_capture_plan_path),
        "diagram_plan_path": str(diagram_plan_path),
        "video_plan_path": str(video_plan_path),
        "reference_template_cleanup_path": str(reference_template_cleanup_path),
        "submission_package_path": str(submission_package_path),
        "insert_config_path": str(insert_config_path),
        "images_dir": str(images_dir),
        "template_manifest_path": str(template_manifest_path),
        "task_scripts_dir": str(task_scripts_dir),
        "browser_capture_script": str((Path(__file__).resolve().parent / "capture_frontend_screenshots.py").resolve()),
        "diagram_assets_script": str((Path(__file__).resolve().parent / "generate_diagram_assets.py").resolve()),
        "video_process_script": str((Path(__file__).resolve().parent / "video_process.py").resolve()),
        "blank_template_script": str((Path(__file__).resolve().parent / "prepare_blank_template.py").resolve()),
        "submission_package_script": str((Path(__file__).resolve().parent / "package_submission.py").resolve()),
        "image_concurrency_report_path": str((output_dir / "image_concurrency_report.json").resolve()),
        "delivery_review_path": str(delivery_review_path),
        "docx_scripts": {
            "fill": str(fill_script_path),
            "insert": str(insert_script_path),
            "verify": str(verify_script_path),
        },
        "commands": {
            "fill_command": 'python "{fill_script}" --template "{template}" --copywriting "{copywriting}" --output "{output_docx}"',
            "insert_command": 'python "{insert_script}" --docx "{output_docx}" --insert-config "{insert_config}"',
            "verify_command": 'python "{verify_script}" --docx "{output_docx}" --insert-config "{insert_config}"',
        },
    }

    requirement_checklist = {
        "has_grading_rubric": True,
        "target_tier": "excellent",
        "run_mode": "planning_only",
        "pre_task_required": False,
        "images_required": False,
        "ai_images_required": False,
        "browser_capture_required": False,
        "diagram_assets_required": False,
        "video_required": False,
        "reference_template_cleanup_required": False,
        "submission_package_required": False,
        "ai_visual_review_completed": False,
        "diagram_visual_review_completed": False,
        "video_review_completed": False,
        "allow_zero_images": True,
        "minimum_image_count": 0,
        "image_route_policy": {
            "ai_simulated_enabled": True,
            "browser_capture_enabled": True,
            "diagram_assets_enabled": True,
            "video_processing_enabled": True,
            "ai_simulated_scope": [
                "terminal",
                "command_output",
                "software_configuration"
            ],
            "browser_capture_scope": [
                "local_frontend_page",
                "self_built_app_or_web_practice"
            ],
            "diagram_assets_scope": [
                "function_diagram",
                "flowchart",
                "data_flow_diagram",
                "er_diagram"
            ]
            ,
            "video_processing_scope": [
                "video_analysis",
                "screen_recording",
                "operation_recording"
            ]
        },
        "planned_figures": [],
        "notes": [
            "This file starts in a neutral planning state.",
            "Use run_mode as a descriptive execution label, not as the source of truth for semantics.",
            "Requirement-dependent meaning is agent-decided, not script-decided. Read docs/prompts/prompt_driven_decisions.md before filling this checklist.",
            "If the requirement depends on a pre-task such as building a system, implementing pages, preparing data, or generating intermediate assets, mark pre_task_required=true first.",
            "Pre-task outputs must be recorded in pre_task_plan.json before report writing.",
            "AI screenshots are for terminal, command output, and configuration content.",
            "Browser screenshots are for local frontend pages and self-built app/web practice screenshots.",
            "Diagram assets are for function diagrams, flowcharts, data flow diagrams, and ER diagrams.",
            "If AI images are enabled, do not mark ai_visual_review_completed=true until an agent has visually checked readability, realism, and absence of localhost or distortion.",
            "If diagram assets are enabled, do not mark diagram_visual_review_completed=true until an agent has checked line routing, spacing, and readability.",
            "If video evidence is enabled, fill video_plan.json and inspect the produced analysis/recording before marking video_review_completed=true.",
            "If the requirement asks for a submission package, derive the needed files from the prompt/requirement first, then write submission_package.json and produce submit.zip.",
            "For structural DOCX operations, use vendor/minimax-docx first; python-docx is a fallback only for simple fill/cleanup or when minimax-docx is unavailable."
        ]
    }

    requirement_analysis = {
        "status": "needs_agent_analysis",
        "source_of_truth": [
            "requirement document",
            "template structure",
            "project artifacts"
        ],
        "decision_summary": "",
        "pre_task_judgment": {
            "required": None,
            "reason": "",
            "expected_outputs": []
        },
        "route_judgment": {
            "ai_simulated": {
                "needed": None,
                "reason": ""
            },
            "browser_capture": {
                "needed": None,
                "reason": ""
            },
            "diagram_assets": {
                "needed": None,
                "reason": ""
            },
            "video_analysis": {
                "needed": None,
                "reason": ""
            },
            "screen_recording": {
                "needed": None,
                "reason": ""
            }
        },
        "figure_strategy": {
            "images_required": None,
            "minimum_image_count": None,
            "why": "",
            "planned_anchors": []
        },
        "template_strategy": {
            "fill_mode": "",
            "cleanup_mode": "",
            "why": "",
            "notes": []
        },
        "submission_strategy": {
            "required": None,
            "archive_name": "submit.zip",
            "deliverables": [],
            "why": ""
        },
        "notes": [
            "This file should hold the agent's actual reasoning outcomes after reading the requirement.",
            "Scripts may execute and validate later, but they should not replace this analysis."
        ]
    }

    pre_task_plan = {
        "enabled": False,
        "completed": False,
        "task_type": "",
        "objective": "",
        "requirements_dependency": "",
        "execution_summary": "",
        "output_summary": "",
        "output_paths": [],
        "output_artifacts": [],
        "report_usage_notes": [
            "If the requirement first asks for a system/app/page/database/demo result, complete that pre-task before writing the report.",
            "Record what was built, what files/artifacts were produced, and which later report sections depend on those outputs."
        ]
    }

    prompt_config = {
        "total_count": 0,
        "resolution": "2560x1440",
        "output_dir": str(images_dir),
        "max_workers": 1,
        "max_retries": 3,
        "retry_delay": 2,
        "concurrency_source": "bootstrap_default_until_agent_runs_test_image_concurrency",
        "concurrency_report": str((output_dir / "image_concurrency_report.json").resolve()),
        "image_policy": {
            "default_mode": "screenshot_strict",
            "auto_append_negative": True,
            "fail_on_prompt_risk": True,
            "forbidden_terms": [
                "流程图",
                "架构图",
                "讲解板",
                "说明面板",
                "悬浮标注",
                "箭头标注",
                "海报",
                "AI生成",
                "示意图",
                "poster",
                "callout",
                "annotation",
                "flowchart",
                "diagram"
            ],
            "ui_density": "low_information_density",
            "crop_browser_chrome": True,
            "forbid_localhost_or_dev_url": True
        },
        "global_prompt": "",
        "images": []
    }

    browser_capture_plan = {
        "enabled": False,
        "capability_check_required": True,
        "startup_command": "",
        "startup_cwd": "",
        "base_url": "",
        "target_kind": "frontend_or_self_built_app",
        "screenshots": [],
        "notes": [
            "Use only for local frontend pages or self-built app/web practice screenshots.",
            "Do not use this route for terminal or software configuration figures.",
            "The actual pages, startup command, and base URL must come from the real project requirement and local app, not from a generic default."
        ]
    }

    diagram_plan = {
        "enabled": False,
        "generator": "python_pil",
        "diagrams": [],
        "notes": [
            "Use only for function diagrams, flowcharts, data flow diagrams, and ER diagrams.",
            "Do not use this route for terminal screenshots or local frontend page screenshots.",
            "Diagram content, labels, and structure must be derived from the assignment and project, not from a canned template."
        ]
    }

    video_plan = {
        "enabled": False,
        "backend_priority": [
            "pyav",
            "opencv",
            "ffprobe"
        ],
        "recording_priority": [
            "mss_opencv",
            "obs_websocket_or_cli",
            "ffmpeg_gdigrab"
        ],
        "videos": [],
        "recordings": [],
        "notes": [
            "Use PyAV first for precise FFmpeg-backed video parsing and frame sampling.",
            "If PyAV is unavailable, fall back to OpenCV or ffprobe metadata.",
            "Use mss + OpenCV for short local screen evidence recordings.",
            "If a real recording-software workflow is required, use OBS WebSocket/CLI as the external fallback and record the exact command.",
            "Whether video is required at all must come from the assignment requirement, not from the existence of this module."
        ]
    }

    reference_template_cleanup = {
        "enabled": False,
        "filled_reference_docx": "",
        "blank_template_docx": "",
        "keep_cover_until_first_level_1_heading": True,
        "keep_heading_levels": [
            1,
            2
        ],
        "remove_after_first_level_1_heading": [
            "Body Text paragraphs",
            "List Paragraph items such as (1), (2), 1., and 1)",
            "Normal explanatory paragraphs including figure/table captions",
            "all body tables",
            "all body images"
        ],
        "must_preserve": [
            "cover area exactly before the first level-1 heading",
            "cover tables and fixed text",
            "cover spacing",
            "TOC entry zone before the first level-1 heading",
            "level-1 and level-2 heading text"
        ],
        "minimax_docx_policy": "For structural DOCX operations, read and prefer vendor/minimax-docx. Use python-docx only for this simple cleanup script or as a documented fallback.",
        "fallback_reason": ""
    }

    submission_package = {
        "enabled": False,
        "source_root": str(output_dir),
        "output_zip": str(output_dir / "submit.zip"),
        "include_paths": [],
        "exclude_globs": [
            "*.tmp",
            "~$*",
            "*.log"
        ],
        "flatten": False,
        "allow_custom_zip_name": False,
        "prompt_driven_rule": "Do not guess blindly. Read the requirement/prompt, decide the exact files that need to be submitted, then update include_paths so the final archive is named submit.zip.",
        "notes": [
            "Use prompt-driven packaging: derive what to submit from the assignment requirement, not from a fixed global archive recipe.",
            "Default output name must be submit.zip unless the user explicitly requires a different archive name.",
            "Do not add extra suffixes like 'AI版', '完整版', 'final' to the zip name.",
            "Also produce a submit/ folder alongside submit.zip for easy inspection and modification.",
            "Do not zip the whole workspace by habit. Include only what the prompt/requirement actually asks for.",
            "Naming must follow the source document requirements; do not add suffixes when none are specified."
        ]
    }

    insert_config = {
        "target_docx": str(output_docx),
        "images": {},
        "placements": {}
    }

    copywriting = """# Copywriting

> 先完成 `requirement_analysis.json`，再更新 `requirement_checklist.json`。
> 初始化默认是 planning_only，不要直接 run。
> 如果需求本身依赖预任务，例如先做系统、先搭页面、先生成数据库设计或先准备实验产物，先完成 `pre_task_plan.json`。
> 预任务完成后，再结合预任务输出与原始需求来写报告正文。
> 不要套固定路线名，由 agent 根据需求判断这次 run 需要哪些路线组合。
> `ai_simulated` 只负责终端、命令输出、软件配置。
> `browser_capture` 只负责本地前端页面和自己编写的 app/web 实操截图。
> `diagram_assets` 只负责功能图、流程图、数据流图、ER 图。
> 如需视频证据，再决定使用 `video_analysis`、`screen_recording` 或两者结合。
> 每个 `{{img_XX}}` 占位符都应该放在图前引导句与图后分析段之间。
"""

    write_script_stub(fill_script_path, "fill")
    write_script_stub(insert_script_path, "insert")
    write_script_stub(verify_script_path, "verify")

    workflow_path.write_text(json.dumps(workflow, ensure_ascii=False, indent=2), encoding="utf-8")
    checklist_path.write_text(json.dumps(requirement_checklist, ensure_ascii=False, indent=2), encoding="utf-8")
    requirement_analysis_path.write_text(json.dumps(requirement_analysis, ensure_ascii=False, indent=2), encoding="utf-8")
    pre_task_plan_path.write_text(json.dumps(pre_task_plan, ensure_ascii=False, indent=2), encoding="utf-8")
    prompt_config_path.write_text(json.dumps(prompt_config, ensure_ascii=False, indent=2), encoding="utf-8")
    browser_capture_plan_path.write_text(json.dumps(browser_capture_plan, ensure_ascii=False, indent=2), encoding="utf-8")
    diagram_plan_path.write_text(json.dumps(diagram_plan, ensure_ascii=False, indent=2), encoding="utf-8")
    video_plan_path.write_text(json.dumps(video_plan, ensure_ascii=False, indent=2), encoding="utf-8")
    reference_template_cleanup_path.write_text(json.dumps(reference_template_cleanup, ensure_ascii=False, indent=2), encoding="utf-8")
    submission_package_path.write_text(json.dumps(submission_package, ensure_ascii=False, indent=2), encoding="utf-8")
    insert_config_path.write_text(json.dumps(insert_config, ensure_ascii=False, indent=2), encoding="utf-8")
    copywriting_path.write_text(copywriting, encoding="utf-8")
    template_manifest_path.write_text(json.dumps(template_manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"Initialized run directory: {output_dir}")
    print(f"  workflow: {workflow_path}")
    print(f"  requirement checklist: {checklist_path}")
    print(f"  requirement analysis: {requirement_analysis_path}")
    print(f"  pre-task plan: {pre_task_plan_path}")
    print(f"  copywriting: {copywriting_path}")
    print(f"  prompt config: {prompt_config_path}")
    print(f"  browser capture plan: {browser_capture_plan_path}")
    print(f"  diagram plan: {diagram_plan_path}")
    print(f"  video plan: {video_plan_path}")
    print(f"  reference template cleanup: {reference_template_cleanup_path}")
    print(f"  submission package: {submission_package_path}")
    print(f"  insert config: {insert_config_path}")
    print(f"  template manifest: {template_manifest_path}")
    print(f"  task scripts: {task_scripts_dir}")


if __name__ == "__main__":
    main()
