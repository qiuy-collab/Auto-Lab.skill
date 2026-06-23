import argparse
import json
import os
import re
import subprocess
import sys
from pathlib import Path


PLACEHOLDER_RE = re.compile(r"\{\{(img_\d{2})\}\}")


def parse_args():
    parser = argparse.ArgumentParser(description="Run and validate the auto-lab workflow.")
    subparsers = parser.add_subparsers(dest="command", required=True)
    for name in ("validate", "gate", "images", "video", "package", "run"):
        sub = subparsers.add_parser(name)
        sub.add_argument("--workflow", required=True, help="Path to workflow.json")
    gate_parser = subparsers.choices.get("gate")
    if gate_parser:
        gate_parser.add_argument("--phase", choices=["1","2","3","4","5","6","all"], default="all",
                                 help="Run a single phase gate (1-6) or all. Default: all")
    return parser.parse_args()


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def save_json(path: Path, data):
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def load_workflow(path_str: str):
    path = Path(path_str).expanduser().resolve()
    workflow = load_json(path)
    workflow["_workflow_path"] = str(path)
    if "pre_task_plan_path" not in workflow and workflow.get("output_dir"):
        workflow["pre_task_plan_path"] = str((Path(workflow["output_dir"]) / "pre_task_plan.json").resolve())
    if "requirement_analysis_path" not in workflow and workflow.get("output_dir"):
        workflow["requirement_analysis_path"] = str((Path(workflow["output_dir"]) / "requirement_analysis.json").resolve())
    if "video_plan_path" not in workflow and workflow.get("output_dir"):
        workflow["video_plan_path"] = str((Path(workflow["output_dir"]) / "video_plan.json").resolve())
    if "reference_template_cleanup_path" not in workflow and workflow.get("output_dir"):
        workflow["reference_template_cleanup_path"] = str((Path(workflow["output_dir"]) / "reference_template_cleanup.json").resolve())
    if "submission_package_path" not in workflow and workflow.get("output_dir"):
        workflow["submission_package_path"] = str((Path(workflow["output_dir"]) / "submission_package.json").resolve())
    root = Path(__file__).resolve().parent
    workflow.setdefault("video_process_script", str((root / "video_process.py").resolve()))
    workflow.setdefault("blank_template_script", str((root / "prepare_blank_template.py").resolve()))
    workflow.setdefault("submission_package_script", str((root / "package_submission.py").resolve()))
    return workflow


def placeholder_keys(copywriting_path: Path):
    return sorted(set(PLACEHOLDER_RE.findall(copywriting_path.read_text(encoding="utf-8"))))


def lint_prompt_config(prompt_config):
    policy = prompt_config.get("image_policy", {})
    forbidden_terms = [term.lower() for term in policy.get("forbidden_terms", [])]
    default_mode = policy.get("default_mode", "screenshot_strict")
    errors = []
    for image in prompt_config.get("images", []):
        mode = image.get("mode", default_mode)
        prompt = image.get("prompt", "")
        lower_prompt = prompt.lower()
        if mode == "screenshot_strict":
            bad_terms = [term for term in forbidden_terms if term in lower_prompt]
            if bad_terms:
                errors.append(f"{image.get('name', '<unknown>')}: forbidden terms for screenshot mode: {', '.join(bad_terms)}")
    return errors


def browser_capture_runtime_available():
    try:
        import importlib

        importlib.import_module("playwright.sync_api")
    except Exception:
        return False
    browser_candidates = [
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"
    ]
    return any(os.path.exists(path) for path in browser_candidates)


def validate_requirement_checklist(checklist, output):
    checklist.setdefault("pre_task_required", False)
    checklist.setdefault("video_required", False)
    checklist.setdefault("reference_template_cleanup_required", False)
    checklist.setdefault("submission_package_required", False)
    checklist.setdefault("video_review_completed", False)
    required_keys = [
        "has_grading_rubric",
        "target_tier",
        "run_mode",
        "pre_task_required",
        "images_required",
        "ai_images_required",
        "browser_capture_required",
        "diagram_assets_required",
        "video_required",
        "reference_template_cleanup_required",
        "submission_package_required",
        "ai_visual_review_completed",
        "diagram_visual_review_completed",
        "video_review_completed",
        "allow_zero_images",
        "minimum_image_count",
        "planned_figures",
        "image_route_policy"
    ]
    missing = [key for key in required_keys if key not in checklist]
    if missing:
        raise SystemExit("requirement_checklist.json is missing keys:\n" + "\n".join(missing))
    output.append(f"Checklist target tier: {checklist.get('target_tier')}")
    output.append(f"Checklist run mode: {checklist.get('run_mode')}")
    output.append(f"Checklist pre-task required: {checklist.get('pre_task_required')}")
    output.append(f"Checklist images required: {checklist.get('images_required')}")
    output.append(f"Checklist AI images required: {checklist.get('ai_images_required')}")
    output.append(f"Checklist browser capture required: {checklist.get('browser_capture_required')}")
    output.append(f"Checklist diagram assets required: {checklist.get('diagram_assets_required')}")
    output.append(f"Checklist video required: {checklist.get('video_required')}")
    output.append(f"Checklist reference template cleanup required: {checklist.get('reference_template_cleanup_required')}")
    output.append(f"Checklist submission package required: {checklist.get('submission_package_required')}")
    output.append(f"Checklist AI visual review completed: {checklist.get('ai_visual_review_completed')}")
    output.append(f"Checklist diagram visual review completed: {checklist.get('diagram_visual_review_completed')}")
    output.append(f"Checklist video review completed: {checklist.get('video_review_completed')}")


def validate_requirement_analysis(checklist, analysis, output):
    required_keys = [
        "status",
        "source_of_truth",
        "decision_summary",
        "pre_task_judgment",
        "route_judgment",
        "figure_strategy",
        "template_strategy",
        "submission_strategy",
    ]
    missing = [key for key in required_keys if key not in analysis]
    if missing:
        raise SystemExit("requirement_analysis.json is missing keys:\n" + "\n".join(missing))

    output.append(f"Requirement analysis status: {analysis.get('status')}")
    decision_summary = str(analysis.get("decision_summary", "")).strip()
    any_execution_flags = any(
        bool(checklist.get(key, False))
        for key in (
            "pre_task_required",
            "images_required",
            "ai_images_required",
            "browser_capture_required",
            "diagram_assets_required",
            "video_required",
            "reference_template_cleanup_required",
            "submission_package_required",
        )
    )
    if any_execution_flags and not decision_summary:
        raise SystemExit(
            "requirement_analysis.json.decision_summary is empty, but the checklist already enables requirement-dependent execution flags"
        )


def validate_pre_task_plan(checklist, pre_task_plan):
    pre_task_required = bool(checklist.get("pre_task_required", False))
    if not pre_task_required:
        return

    if not pre_task_plan.get("enabled", False):
        raise SystemExit("Checklist requires a pre-task, but pre_task_plan.json is not enabled")
    if not pre_task_plan.get("completed", False):
        raise SystemExit("Checklist requires a pre-task, but pre_task_plan.json is not marked completed")
    if not str(pre_task_plan.get("task_type", "")).strip():
        raise SystemExit("Checklist requires a pre-task, but pre_task_plan.json.task_type is empty")
    if not str(pre_task_plan.get("objective", "")).strip():
        raise SystemExit("Checklist requires a pre-task, but pre_task_plan.json.objective is empty")
    if not str(pre_task_plan.get("output_summary", "")).strip():
        raise SystemExit("Checklist requires a pre-task, but pre_task_plan.json.output_summary is empty")

    output_paths = pre_task_plan.get("output_paths", [])
    output_artifacts = pre_task_plan.get("output_artifacts", [])
    if not output_paths and not output_artifacts:
        raise SystemExit("Checklist requires a pre-task, but pre_task_plan.json has no output_paths or output_artifacts")


def default_pre_task_plan():
    return {
        "enabled": False,
        "completed": False,
        "task_type": "",
        "objective": "",
        "requirements_dependency": "",
        "execution_summary": "",
        "output_summary": "",
        "output_paths": [],
        "output_artifacts": [],
        "report_usage_notes": []
    }


def default_video_plan():
    return {"enabled": False, "videos": [], "recordings": []}


def default_reference_template_cleanup():
    return {"enabled": False, "filled_reference_docx": "", "blank_template_docx": ""}


def default_submission_package():
    return {"enabled": False, "include_paths": [], "output_zip": "submit.zip"}


def validate_video_plan(checklist, video_plan):
    if not checklist.get("video_required", False):
        return
    if not video_plan.get("enabled", False):
        raise SystemExit("Checklist requires video processing, but video_plan.json is not enabled")
    if not video_plan.get("videos") and not video_plan.get("recordings"):
        raise SystemExit("Checklist requires video processing, but video_plan.json has no videos or recordings")
    if not checklist.get("video_review_completed", False):
        raise SystemExit("Checklist requires video processing, but video_review_completed is not true")


def validate_reference_template_cleanup(checklist, cleanup_plan):
    if not checklist.get("reference_template_cleanup_required", False):
        return
    if not cleanup_plan.get("enabled", False):
        raise SystemExit("Checklist requires reference template cleanup, but reference_template_cleanup.json is not enabled")
    if not str(cleanup_plan.get("filled_reference_docx", "")).strip():
        raise SystemExit("reference_template_cleanup.json.filled_reference_docx is empty")
    if not str(cleanup_plan.get("blank_template_docx", "")).strip():
        raise SystemExit("reference_template_cleanup.json.blank_template_docx is empty")


def validate_submission_package(checklist, package_plan):
    if not checklist.get("submission_package_required", False):
        return
    if not package_plan.get("enabled", False):
        raise SystemExit("Checklist requires a submission package, but submission_package.json is not enabled")
    if not str(package_plan.get("output_zip", "")).strip():
        raise SystemExit("submission_package.json.output_zip is empty")
    if Path(str(package_plan.get("output_zip"))).name.lower() != "submit.zip":
        raise SystemExit("submission package output must be named submit.zip unless the user explicitly changed the rule")
    if not package_plan.get("include_paths"):
        raise SystemExit("submission_package.json.include_paths is empty")
    for index, item in enumerate(package_plan.get("include_paths", []), start=1):
        if not isinstance(item, dict):
            raise SystemExit(f"submission_package.json.include_paths[{index}] must be an object with at least a path field")
        if not str(item.get("path", "")).strip():
            raise SystemExit(f"submission_package.json.include_paths[{index}].path is empty")


def validate_route_boundaries(checklist, planned_figures, browser_capture_plan, diagram_plan):
    route_policy = checklist.get("image_route_policy", {})
    ai_scopes = set(route_policy.get("ai_simulated_scope", []))
    browser_scopes = set(route_policy.get("browser_capture_scope", []))
    diagram_scopes = set(route_policy.get("diagram_assets_scope", []))
    run_mode = checklist.get("run_mode")

    if not str(run_mode).strip():
        raise SystemExit("requirement_checklist.json.run_mode must be a non-empty descriptive label")

    planned_by_name = {}
    for figure in planned_figures:
        name = figure.get("name")
        source_mode = figure.get("source_mode")
        scope = figure.get("scope")
        if not name or not source_mode or not scope:
            raise SystemExit("Each planned_figures entry must include name, source_mode, and scope")
        if source_mode == "ai_simulated" and scope not in ai_scopes:
            raise SystemExit(f"{name} is routed to ai_simulated but scope '{scope}' is not allowed for AI images")
        if source_mode == "browser_capture" and scope not in browser_scopes:
            raise SystemExit(f"{name} is routed to browser_capture but scope '{scope}' is not allowed for browser capture")
        if source_mode == "diagram_assets" and scope not in diagram_scopes:
            raise SystemExit(f"{name} is routed to diagram_assets but scope '{scope}' is not allowed for diagram assets")
        if source_mode not in {"ai_simulated", "browser_capture", "diagram_assets"}:
            raise SystemExit(f"{name} has unsupported source_mode: {source_mode}")
        planned_by_name[name] = figure

    for shot in browser_capture_plan.get("screenshots", []):
        name = shot.get("name")
        scope = shot.get("scope")
        if not name:
            raise SystemExit("Each browser_capture_plan screenshot must include a name")
        if name not in planned_by_name:
            raise SystemExit(f"browser_capture_plan screenshot '{name}' is not declared in planned_figures")
        planned_figure = planned_by_name[name]
        if planned_figure.get("source_mode") != "browser_capture":
            raise SystemExit(f"browser_capture_plan screenshot '{name}' is not routed as browser_capture in planned_figures")
        if scope and scope not in browser_scopes:
            raise SystemExit(f"browser_capture_plan screenshot '{name}' has invalid browser scope '{scope}'")

    for diagram in diagram_plan.get("diagrams", []):
        name = diagram.get("name")
        kind = diagram.get("kind")
        if not name:
            raise SystemExit("Each diagram_plan entry must include a name")
        if name not in planned_by_name:
            raise SystemExit(f"diagram_plan entry '{name}' is not declared in planned_figures")
        planned_figure = planned_by_name[name]
        if planned_figure.get("source_mode") != "diagram_assets":
            raise SystemExit(f"diagram_plan entry '{name}' is not routed as diagram_assets in planned_figures")
        if kind and kind not in diagram_scopes:
            raise SystemExit(f"diagram_plan entry '{name}' has invalid kind '{kind}'")


def validate_workflow(workflow):
    output = []
    requirements = Path(workflow["requirements_path"])
    template = Path(workflow["template_path"])
    output_dir = Path(workflow["output_dir"])
    checklist_path = Path(workflow["requirement_checklist_path"])
    requirement_analysis_path = Path(workflow["requirement_analysis_path"])
    pre_task_plan_path = Path(workflow["pre_task_plan_path"])
    copywriting_path = Path(workflow["copywriting_path"])
    prompt_config_path = Path(workflow["prompt_config_path"])
    browser_capture_plan_path = Path(workflow["browser_capture_plan_path"])
    diagram_plan_path = Path(workflow["diagram_plan_path"])
    video_plan_path = Path(workflow["video_plan_path"])
    reference_template_cleanup_path = Path(workflow["reference_template_cleanup_path"])
    submission_package_path = Path(workflow["submission_package_path"])
    insert_config_path = Path(workflow["insert_config_path"])
    template_manifest_path = Path(workflow["template_manifest_path"])
    docx_scripts = workflow.get("docx_scripts", {})
    fill_script = Path(docx_scripts.get("fill", ""))
    insert_script = Path(docx_scripts.get("insert", ""))
    verify_script = Path(docx_scripts.get("verify", ""))

    missing = [
        str(path)
        for path in (
            requirements,
            template,
            output_dir,
            checklist_path,
            requirement_analysis_path,
            copywriting_path,
            prompt_config_path,
            browser_capture_plan_path,
            diagram_plan_path,
            insert_config_path,
            template_manifest_path,
            fill_script,
            insert_script,
            verify_script
        )
        if not path.exists()
    ]
    if missing:
        raise SystemExit("Missing required files:\n" + "\n".join(missing))

    checklist = load_json(checklist_path)
    validate_requirement_checklist(checklist, output)
    requirement_analysis = load_json(requirement_analysis_path)
    validate_requirement_analysis(checklist, requirement_analysis, output)
    pre_task_plan = load_json(pre_task_plan_path) if pre_task_plan_path.exists() else default_pre_task_plan()
    validate_pre_task_plan(checklist, pre_task_plan)
    prompt_config = load_json(prompt_config_path)
    browser_capture_plan = load_json(browser_capture_plan_path)
    diagram_plan = load_json(diagram_plan_path)
    video_plan = load_json(video_plan_path) if video_plan_path.exists() else default_video_plan()
    reference_template_cleanup = load_json(reference_template_cleanup_path) if reference_template_cleanup_path.exists() else default_reference_template_cleanup()
    submission_package = load_json(submission_package_path) if submission_package_path.exists() else default_submission_package()
    insert_config = load_json(insert_config_path)
    validate_video_plan(checklist, video_plan)
    validate_reference_template_cleanup(checklist, reference_template_cleanup)
    validate_submission_package(checklist, submission_package)

    copy_keys = placeholder_keys(copywriting_path)
    prompt_keys = [item["name"] for item in prompt_config.get("images", [])]
    insert_keys = sorted(insert_config.get("images", {}).keys())
    planned_figures = checklist.get("planned_figures", [])
    validate_route_boundaries(checklist, planned_figures, browser_capture_plan, diagram_plan)

    output.append(f"Pre-task enabled: {pre_task_plan.get('enabled')}")
    output.append(f"Pre-task completed: {pre_task_plan.get('completed')}")
    output.append(f"Placeholders in copywriting.md: {len(copy_keys)}")
    output.append(f"Images in prompt_config.json: {len(prompt_keys)}")
    output.append(f"Images in insert_config.json: {len(insert_keys)}")
    output.append(f"Template manifest: {template_manifest_path.name}")
    output.append(f"Video plan enabled: {video_plan.get('enabled')}")
    output.append(f"Reference template cleanup enabled: {reference_template_cleanup.get('enabled')}")
    output.append(f"Submission package enabled: {submission_package.get('enabled')}")

    images_required = bool(checklist.get("images_required", False))
    ai_images_required = bool(checklist.get("ai_images_required", False))
    browser_capture_required = bool(checklist.get("browser_capture_required", False))
    diagram_assets_required = bool(checklist.get("diagram_assets_required", False))
    allow_zero_images = bool(checklist.get("allow_zero_images", False))
    minimum_image_count = int(checklist.get("minimum_image_count", 0))
    planning_only = str(checklist.get("run_mode", "")).strip().lower() == "planning_only"

    ai_planned = sorted([fig["name"] for fig in planned_figures if fig.get("source_mode") == "ai_simulated"])
    browser_planned = sorted([fig["name"] for fig in planned_figures if fig.get("source_mode") == "browser_capture"])
    diagram_planned = sorted([fig["name"] for fig in planned_figures if fig.get("source_mode") == "diagram_assets"])
    browser_configured = sorted([item.get("name") for item in browser_capture_plan.get("screenshots", []) if item.get("name")])
    diagram_configured = sorted([item.get("name") for item in diagram_plan.get("diagrams", []) if item.get("name")])

    if ai_images_required and sorted(prompt_keys) != ai_planned:
        raise SystemExit("AI prompt image names do not match planned_figures entries for source_mode=ai_simulated")
    if browser_capture_required and browser_configured != browser_planned:
        raise SystemExit("browser_capture_plan screenshot names do not match planned_figures entries for source_mode=browser_capture")
    if diagram_assets_required and diagram_configured != diagram_planned:
        raise SystemExit("diagram_plan names do not match planned_figures entries for source_mode=diagram_assets")

    if insert_keys and not set(prompt_keys).issubset(set(insert_keys)):
        raise SystemExit("insert_config.json is missing one or more AI image keys")

    if not planning_only and images_required:
        if minimum_image_count <= 0 and not allow_zero_images:
            raise SystemExit("images_required is true, but minimum_image_count is not positive")
        if len(planned_figures) < minimum_image_count:
            raise SystemExit(
                f"Planned figure count below checklist minimum. Need at least {minimum_image_count}, found {len(planned_figures)}"
            )
        if len(copy_keys) == 0 and all(not fig.get("heading") for fig in planned_figures):
            raise SystemExit("Images are required, but copywriting.md has no placeholders and planned_figures has no non-empty headings/anchors")

    if browser_capture_required:
        if not browser_capture_runtime_available():
            raise SystemExit("Checklist requires browser capture, but local browser screenshot runtime is not available")
        if not browser_capture_plan.get("enabled", False):
            raise SystemExit("Checklist requires browser capture, but browser_capture_plan.json is not enabled")
        if not browser_capture_plan.get("startup_command"):
            raise SystemExit("Checklist requires browser capture, but startup_command is empty")
        if not browser_capture_plan.get("base_url"):
            raise SystemExit("Checklist requires browser capture, but base_url is empty")
        if not browser_capture_plan.get("screenshots"):
            raise SystemExit("Checklist requires browser capture, but screenshots is empty")

    if diagram_assets_required:
        if not diagram_plan.get("enabled", False):
            raise SystemExit("Checklist requires diagram assets, but diagram_plan.json is not enabled")
        if not diagram_plan.get("diagrams"):
            raise SystemExit("Checklist requires diagram assets, but diagram_plan.json has no diagrams")

    if ai_images_required and not bool(checklist.get("ai_visual_review_completed", False)):
        raise SystemExit("Checklist requires AI images, but ai_visual_review_completed is not true")

    if ai_images_required:
        max_workers = int(prompt_config.get("max_workers", 0))
        if max_workers <= 0:
            raise SystemExit("prompt_config.json.max_workers must be a positive integer when AI images are required")

    if diagram_assets_required and not bool(checklist.get("diagram_visual_review_completed", False)):
        raise SystemExit("Checklist requires diagram assets, but diagram_visual_review_completed is not true")

    prompt_errors = lint_prompt_config(prompt_config)
    if prompt_errors and prompt_config.get("image_policy", {}).get("fail_on_prompt_risk", True):
        raise SystemExit("Prompt policy validation failed:\n" + "\n".join(prompt_errors))

    return checklist, requirement_analysis, prompt_config, browser_capture_plan, diagram_plan, video_plan, reference_template_cleanup, submission_package, insert_config, output


def ensure_docx_scripts_customized(workflow):
    stubbed = []
    for name in ("fill", "insert", "verify"):
        script_path = Path(workflow["docx_scripts"][name])
        if "AUTO_LAB_TEMPLATE_SCRIPT_STUB = True" in script_path.read_text(encoding="utf-8"):
            stubbed.append(f"{name}: {script_path}")
    if stubbed:
        raise SystemExit(
            "Template-specific docx scripts are still stubs. Customize them first:\n" + "\n".join(stubbed)
        )


def update_insert_config_from_prompts(workflow, prompt_config, insert_config):
    images_dir = Path(workflow["images_dir"])
    insert_config["target_docx"] = workflow["output_docx"]
    current_placements = insert_config.get("placements", {})
    current_images = insert_config.get("images", {})
    for item in prompt_config.get("images", []):
        current_images[item["name"]] = str((images_dir / f"{item['name']}.png").resolve())
        current_placements.setdefault(item["name"], {})
    insert_config["images"] = current_images
    insert_config["placements"] = current_placements
    save_json(Path(workflow["insert_config_path"]), insert_config)


def update_insert_config_from_browser_plan(workflow, browser_capture_plan, insert_config):
    images_dir = Path(workflow["images_dir"])
    current_placements = insert_config.get("placements", {})
    current_images = insert_config.get("images", {})
    for item in browser_capture_plan.get("screenshots", []):
        current_images[item["name"]] = str((images_dir / f"{item['name']}.png").resolve())
        current_placements.setdefault(item["name"], {})
    insert_config["images"] = current_images
    insert_config["placements"] = current_placements
    save_json(Path(workflow["insert_config_path"]), insert_config)


def update_insert_config_from_diagram_plan(workflow, diagram_plan, insert_config):
    images_dir = Path(workflow["images_dir"])
    current_placements = insert_config.get("placements", {})
    current_images = insert_config.get("images", {})
    for item in diagram_plan.get("diagrams", []):
        current_images[item["name"]] = str((images_dir / f"{item['name']}.png").resolve())
        current_placements.setdefault(item["name"], {})
    insert_config["images"] = current_images
    insert_config["placements"] = current_placements
    save_json(Path(workflow["insert_config_path"]), insert_config)


def run_generate_py(workflow):
    root = Path(__file__).resolve().parent
    prompt_config_path = Path(workflow["prompt_config_path"])

    # B7: Test upstream before batch
    print("Testing upstream image generation capability...")
    check_result = subprocess.run(
        [sys.executable, str(root / "generate_images.py"), "--check"],
        capture_output=True, text=True
    )
    if check_result.returncode != 0:
        raise SystemExit(
            "Upstream image generation is not available. "
            "Fix the API connection before running batch generation.\n"
            f"Error: {check_result.stderr or check_result.stdout}"
        )
    print("Upstream check passed, starting batch generation...")

    subprocess.run([sys.executable, str(root / "generate_images.py"), "--config", str(prompt_config_path)], check=True)


def run_browser_capture(workflow):
    script_path = Path(workflow["browser_capture_script"])
    subprocess.run([sys.executable, str(script_path), "--workflow", workflow["_workflow_path"]], check=True)


def run_diagram_assets(workflow):
    script_path = Path(workflow["diagram_assets_script"])
    subprocess.run([sys.executable, str(script_path), "--workflow", workflow["_workflow_path"]], check=True)


def run_video_processing(workflow, video_plan):
    script_path = Path(workflow["video_process_script"])
    output_dir = Path(workflow["output_dir"])
    for item in video_plan.get("videos", []):
        input_path = item.get("input")
        output_path = item.get("analysis_output") or str(output_dir / f"{item.get('name', 'video')}_analysis.json")
        sample_dir = item.get("sample_frames_dir") or str(output_dir / "video_frames" / item.get("name", "video"))
        max_frames = str(item.get("max_frames", 5))
        if not input_path:
            raise SystemExit("Each video_plan.videos entry must include input")
        subprocess.run(
            [
                sys.executable,
                str(script_path),
                "analyze",
                "--input",
                input_path,
                "--output",
                output_path,
                "--sample-frames-dir",
                sample_dir,
                "--max-frames",
                max_frames,
            ],
            check=True,
        )
    for item in video_plan.get("recordings", []):
        output_path = item.get("output")
        if not output_path:
            raise SystemExit("Each video_plan.recordings entry must include output")
        subprocess.run(
            [
                sys.executable,
                str(script_path),
                "record-screen",
                "--output",
                output_path,
                "--seconds",
                str(item.get("seconds", 10)),
                "--fps",
                str(item.get("fps", 15)),
                "--monitor",
                str(item.get("monitor", 1)),
            ],
            check=True,
        )


def run_reference_template_cleanup(workflow, cleanup_plan):
    if not cleanup_plan.get("enabled", False):
        return
    script_path = Path(workflow["blank_template_script"])
    report_path = cleanup_plan.get("report") or str(Path(workflow["output_dir"]) / "reference_template_cleanup_report.json")
    subprocess.run(
        [
            sys.executable,
            str(script_path),
            "--input",
            cleanup_plan["filled_reference_docx"],
            "--output",
            cleanup_plan["blank_template_docx"],
            "--report",
            report_path,
            "--keep-levels",
            ",".join(str(item) for item in cleanup_plan.get("keep_heading_levels", [1, 2])),
        ],
        check=True,
    )


def run_submission_package(workflow):
    script_path = Path(workflow["submission_package_script"])
    subprocess.run(
        [
            sys.executable,
            str(script_path),
            "--config",
            workflow["submission_package_path"],
        ],
        check=True,
    )


def expand_command(template, workflow):
    replacements = {
        "{template}": workflow["template_path"],
        "{requirements}": workflow["requirements_path"],
        "{output_dir}": workflow["output_dir"],
        "{output_docx}": workflow["output_docx"],
        "{copywriting}": workflow["copywriting_path"],
        "{prompt_config}": workflow["prompt_config_path"],
        "{insert_config}": workflow["insert_config_path"],
        "{images_dir}": workflow["images_dir"],
        "{fill_script}": workflow["docx_scripts"]["fill"],
        "{insert_script}": workflow["docx_scripts"]["insert"],
        "{verify_script}": workflow["docx_scripts"]["verify"]
    }
    command = template
    for key, value in replacements.items():
        command = command.replace(key, value)
    return command


def run_manifest_commands(workflow):
    for command_name in ("fill_command", "insert_command", "verify_command"):
        template = workflow.get("commands", {}).get(command_name, "").strip()
        if template:
            command = expand_command(template, workflow)
            print(f"Running {command_name}: {command}")
            subprocess.run(command, check=True, shell=True)


# ══════════════════════════════════════════════════════════════
# Phase Gates — 映射自 SKILL.md 流程图，逐级硬关卡
# 每个 Gate 返回 (pass: bool, errors: list[str], warnings: list[str])
# ══════════════════════════════════════════════════════════════

DEFAULT_README_PATTERNS = [
    "README.md",
    "readme.md",
    "README.txt",
    "STARTUP.md",
    "QUICKSTART.md",
]
FORBIDDEN_IN_OUTPUT = [
    "localhost",
    "127.0.0.1",
    "AI 生成",
    "AI生成",
    "根据模板",
    "请在此处填写",
    "占位",
    "（此处替换",
    "[TODO]",
    "TODO：",
    "示例内容",
    "样例",
    "dashboard",
    "welcome",
    "placeholder",
    "sample data",
    "字号要求",
    "行距要求",
    "格式要求",
    "字体要求",
]
FORBIDDEN_IN_OUTPUT_RE = re.compile(
    "|".join(re.escape(t) for t in FORBIDDEN_IN_OUTPUT),
    re.IGNORECASE,
)

VENDOR_SKILL_FILES = {
    "baseline-ui": "vendor/baseline-ui/SKILL.md",
    "frontend-design": "vendor/frontend-design/SKILL.md",
    "webapp-testing": "vendor/webapp-testing/SKILL.md",
    "minimax-docx": "vendor/minimax-docx/SKILL.md",
}


def _skill_root(workflow_script_dir: Path) -> Path:
    return workflow_script_dir.parent


# ── Gate 1: INIT ────────────────────────────────────────────
# 流程图节点: B (env_check) → C (init_run) → D (agent 已决策)

def gate_init(workflow: dict, output_dir: Path, checklist: dict) -> tuple:
    errors, warnings = [], []
    root = _skill_root(Path(workflow.get("_workflow_path", ".")).parent)

    # B: env_check.ps1 — 检查至少 python 可用
    try:
        subprocess.run([sys.executable, "--version"], capture_output=True, check=True)
    except Exception:
        errors.append("[Gate 1/B] Python 不可用，请先运行 env_check.ps1")

    # B: 检查关键 Python 模块
    for mod, pip_name in [("docx", "python-docx"), ("PIL", "pillow")]:
        try:
            __import__(mod)
        except ImportError:
            errors.append(f"[Gate 1/B] Python 模块 {pip_name} 未安装，请运行 env_check.ps1")

    # C: init_run 产生的目录和文件必须存在
    if not output_dir.exists():
        errors.append(f"[Gate 1/C] 输出目录不存在: {output_dir}，请先运行 init_run.py")
        return (False, errors, warnings)  # 后续检查无意义

    required_init_files = [
        "workflow.json",
        "requirement_checklist.json",
        "requirement_analysis.json",
        "pre_task_plan.json",
        "copywriting.md",
        "prompt_config.json",
        "browser_capture_plan.json",
        "diagram_plan.json",
        "video_plan.json",
        "insert_config.json",
        "template_manifest.json",
    ]
    for fname in required_init_files:
        if not (output_dir / fname).exists():
            errors.append(f"[Gate 1/C] 缺少初始化文件: {fname}，请重新运行 init_run.py")

    task_scripts = ["fill_template.py", "insert_images.py", "verify_template.py"]
    for fname in task_scripts:
        if not (output_dir / "task_scripts" / fname).exists():
            errors.append(f"[Gate 1/C] 缺少任务脚本: task_scripts/{fname}")

    # D: Agent 必须已经做了决策（run_mode 不能还是 planning_only）
    run_mode = str(checklist.get("run_mode", "")).strip().lower()
    if run_mode == "planning_only":
        errors.append(
            "[Gate 1/D] requirement_checklist.json.run_mode 仍为 planning_only，"
            "Agent 必须完成 requirement_analysis.json 并更新 checklist 后才能继续"
        )

    return (len(errors) == 0, errors, warnings)


# ── Gate 2: PRE_TASK ────────────────────────────────────────
# 流程图节点: E (需不需要前置任务) → H (完成前置任务) → I→O (前端子流程) → Q (记录)

def gate_pre_task(workflow: dict, checklist: dict, pre_task_plan: dict) -> tuple:
    errors, warnings = [], []
    output_dir = Path(workflow.get("output_dir", "."))
    pre_task_required = bool(checklist.get("pre_task_required", False))

    if not pre_task_required:
        return (True, [], [])  # 不需要前置任务，Gate 2 直接通过

    # H: pre_task_plan 必须启用且完成
    if not pre_task_plan.get("enabled", False):
        errors.append("[Gate 2/H] checklist 要求前置任务，但 pre_task_plan.json.enabled 仍为 false")
    if not pre_task_plan.get("completed", False):
        errors.append("[Gate 2/H] checklist 要求前置任务，但 pre_task_plan.json.completed 仍为 false")
    if not str(pre_task_plan.get("task_type", "")).strip():
        errors.append("[Gate 2/H] pre_task_plan.json.task_type 为空")
    if not str(pre_task_plan.get("output_summary", "")).strip():
        errors.append("[Gate 2/H] pre_task_plan.json.output_summary 为空")

    output_paths = pre_task_plan.get("output_paths", [])
    output_artifacts = pre_task_plan.get("output_artifacts", [])
    if not output_paths and not output_artifacts:
        errors.append("[Gate 2/H] pre_task_plan.json 没有记录任何 output_paths 或 output_artifacts")

    # 检查 output_paths 中的文件是否真实存在
    missing_outputs = []
    for p in output_paths:
        if not Path(p).exists():
            missing_outputs.append(p)
    if missing_outputs:
        errors.append(
            f"[Gate 2/H] pre_task_plan.json 中 {len(missing_outputs)} 个产出路径不存在: "
            + ", ".join(missing_outputs[:5])
        )

    # I→O: 如果是前端/Web 前置任务，检查 git + vendor skills + README
    task_type = str(pre_task_plan.get("task_type", "")).lower()
    is_frontend = any(kw in task_type for kw in [
        "frontend", "web", "页面", "前端", "app", "ui", "界面"
    ])

    if is_frontend:
        # J: git repo
        project_paths = [Path(p) for p in output_paths if p and Path(p).exists()]
        cwd = Path(workflow.get("_cwd", "."))
        git_found = False
        for candidate in ([cwd] + [p.parent for p in project_paths if p.is_file()] + project_paths):
            if (candidate / ".git").is_dir():
                git_found = True
                break
        if not git_found:
            errors.append(
                "[Gate 2/J] 检测到前端前置任务，但未找到 .git 仓库。"
                "请在实现代码前执行 git init"
            )

        # K→L→O: vendor skills 签署检查
        root = _skill_root(Path(workflow.get("_workflow_path", ".")).parent)
        vendor_missing = []
        for skill_name, rel_path in VENDOR_SKILL_FILES.items():
            if skill_name == "minimax-docx":
                continue  # minimax-docx 在 Gate 5 检查
            skill_file = root / rel_path
            if not skill_file.exists():
                vendor_missing.append(rel_path)
        if vendor_missing:
            errors.append(
                "[Gate 2/K-L-O] 前端前置任务必须使用 vendor skills，但以下文件缺失: "
                + ", ".join(vendor_missing)
            )

        # N: README
        readme_found = False
        for candidate in ([cwd] + [p.parent for p in project_paths if p.is_file()]):
            for pattern in DEFAULT_README_PATTERNS:
                if (candidate / pattern).exists():
                    readme_found = True
                    break
            if readme_found:
                break
        if not readme_found:
            warnings.append(
                "[Gate 2/N] 未找到 README.md/STARTUP.md，建议为前置任务项目编写启动说明"
            )

    return (len(errors) == 0, errors, warnings)


# ── Gate 3: EVIDENCE ────────────────────────────────────────
# 流程图节点: R (分析评分标准) → S (选路线) → T/U/V/W (填配置)

def gate_evidence(workflow: dict, checklist: dict, requirement_analysis: dict) -> tuple:
    errors, warnings = [], []
    output_dir = Path(workflow.get("output_dir", "."))

    # R→S: requirement_analysis 的 decision_summary 不能为空
    decision_summary = str(requirement_analysis.get("decision_summary", "")).strip()
    if not decision_summary:
        errors.append(
            "[Gate 3/R-S] requirement_analysis.json.decision_summary 为空，"
            "Agent 必须分析需求后填写"
        )

    # T: AI 截图路线
    if bool(checklist.get("ai_images_required", False)):
        prompt_config = _load_if_exists(output_dir / "prompt_config.json")
        if not prompt_config:
            errors.append("[Gate 3/T] ai_images_required=true 但 prompt_config.json 不存在")
        else:
            total = int(prompt_config.get("total_count", 0))
            images = prompt_config.get("images", [])
            if total <= 0 and len(images) == 0:
                errors.append(
                    "[Gate 3/T] prompt_config.json 的 total_count 为 0 且 images 为空"
                )
            if not prompt_config.get("global_prompt", "").strip():
                warnings.append("[Gate 3/T] prompt_config.json.global_prompt 为空，建议填写统一环境描述")

    # U: 浏览器截图路线
    if bool(checklist.get("browser_capture_required", False)):
        bc = _load_if_exists(output_dir / "browser_capture_plan.json")
        if not bc:
            errors.append("[Gate 3/U] browser_capture_required=true 但 browser_capture_plan.json 不存在")
        else:
            if not bc.get("enabled", False):
                errors.append("[Gate 3/U] browser_capture_plan.json.enabled 仍为 false")
            if not bc.get("startup_command", "").strip():
                errors.append("[Gate 3/U] browser_capture_plan.json.startup_command 为空")
            if not bc.get("base_url", "").strip():
                errors.append("[Gate 3/U] browser_capture_plan.json.base_url 为空")
            if not bc.get("screenshots"):
                errors.append("[Gate 3/U] browser_capture_plan.json.screenshots 为空")

    # V: 图表路线
    if bool(checklist.get("diagram_assets_required", False)):
        dp = _load_if_exists(output_dir / "diagram_plan.json")
        if not dp:
            errors.append("[Gate 3/V] diagram_assets_required=true 但 diagram_plan.json 不存在")
        else:
            if not dp.get("enabled", False):
                errors.append("[Gate 3/V] diagram_plan.json.enabled 仍为 false")
            if not dp.get("diagrams"):
                errors.append("[Gate 3/V] diagram_plan.json.diagrams 为空")

    # W: 视频路线
    if bool(checklist.get("video_required", False)):
        vp = _load_if_exists(output_dir / "video_plan.json")
        if not vp:
            errors.append("[Gate 3/W] video_required=true 但 video_plan.json 不存在")
        else:
            if not vp.get("enabled", False):
                errors.append("[Gate 3/W] video_plan.json.enabled 仍为 false")
            if not vp.get("videos") and not vp.get("recordings"):
                errors.append("[Gate 3/W] video_plan.json 的 videos 和 recordings 均为空")

    # 交叉检查：copywriting.md 的 placeholder 数量 >= minimum_image_count
    images_required = bool(checklist.get("images_required", False))
    if images_required:
        copywriting_path = output_dir / "copywriting.md"
        if copywriting_path.exists():
            placeholders = placeholder_keys(copywriting_path)
            min_count = int(checklist.get("minimum_image_count", 0))
            if len(placeholders) < min_count:
                errors.append(
                    f"[Gate 3] copywriting.md 中只有 {len(placeholders)} 个占位符，"
                    f"但 checklist 要求至少 {min_count} 张图"
                )
            if len(placeholders) == 0:
                planned = checklist.get("planned_figures", [])
                if not any(fig.get("heading") for fig in planned):
                    errors.append("[Gate 3] copywriting.md 无占位符且 planned_figures 无 heading/锚点")

    # 图片路线总数与 planned_figures 的一致性
    planned = checklist.get("planned_figures", [])
    route_policy = checklist.get("image_route_policy", {})
    for fig in planned:
        name = fig.get("name", "?")
        sm = fig.get("source_mode", "")
        allowed = route_policy.get(f"{sm}_scope", [])
        scope = fig.get("scope", "")
        if allowed and scope not in allowed:
            errors.append(f"[Gate 3] planned_figures 中 '{name}' 的 scope '{scope}' 不在 {sm} 允许范围内")

    return (len(errors) == 0, errors, warnings)


# ── Gate 4: GENERATION ──────────────────────────────────────
# 流程图节点: X (AI 截图) → Y (浏览器截图) → Z (图表) → AA (视频)
#            → AB (视觉审查) → AC (质量门)

def gate_generation(workflow: dict, checklist: dict) -> tuple:
    errors, warnings = [], []
    output_dir = Path(workflow.get("output_dir", "."))
    images_dir = Path(workflow.get("images_dir", str(output_dir / "generated_images")))
    planned = checklist.get("planned_figures", [])

    # X: AI 截图文件必须存在
    if bool(checklist.get("ai_images_required", False)):
        ai_planned = [f["name"] for f in planned if f.get("source_mode") == "ai_simulated"]
        ai_missing = []
        for name in ai_planned:
            png_path = images_dir / f"{name}.png"
            if not png_path.exists():
                ai_missing.append(name)
        if ai_missing:
            errors.append(
                f"[Gate 4/X] {len(ai_missing)} 张 AI 截图未生成: {', '.join(ai_missing[:5])}"
            )
        # AC: 视觉审查标志
        if not bool(checklist.get("ai_visual_review_completed", False)):
            errors.append(
                "[Gate 4/AC] ai_visual_review_completed 仍为 false，"
                "Agent 必须用 docs/prompts/visual_review_rules.md 检查每张 AI 截图"
            )

    # Y: 浏览器截图文件必须存在
    if bool(checklist.get("browser_capture_required", False)):
        bc = _load_if_exists(output_dir / "browser_capture_plan.json") or {}
        for item in bc.get("screenshots", []):
            name = item.get("name", "?")
            png_path = images_dir / f"{name}.png"
            if not png_path.exists():
                errors.append(f"[Gate 4/Y] 浏览器截图未找到: {name} ({png_path})")

    # Z: 图表文件必须存在
    if bool(checklist.get("diagram_assets_required", False)):
        dp = _load_if_exists(output_dir / "diagram_plan.json") or {}
        diagram_missing = []
        for item in dp.get("diagrams", []):
            name = item.get("name", "?")
            png_path = images_dir / f"{name}.png"
            if not png_path.exists():
                diagram_missing.append(name)
        if diagram_missing:
            errors.append(
                f"[Gate 4/Z] {len(diagram_missing)} 张图表未生成: {', '.join(diagram_missing[:5])}"
            )
        # AC: 图表视觉审查
        if not bool(checklist.get("diagram_visual_review_completed", False)):
            errors.append(
                "[Gate 4/AC] diagram_visual_review_completed 仍为 false，"
                "Agent 必须检查布局、路由、可读性和标签无碰撞"
            )

    # AA: 视频产物检查
    if bool(checklist.get("video_required", False)):
        vp = _load_if_exists(output_dir / "video_plan.json") or {}
        for item in vp.get("videos", []):
            output = item.get("analysis_output", "")
            if output and not Path(output).exists():
                errors.append(f"[Gate 4/AA] 视频分析输出不存在: {output}")
        for item in vp.get("recordings", []):
            output = item.get("output", "")
            if output and not Path(output).exists():
                errors.append(f"[Gate 4/AA] 录屏输出不存在: {output}")
        if not bool(checklist.get("video_review_completed", False)):
            errors.append(
                "[Gate 4/AC] video_review_completed 仍为 false"
            )

    return (len(errors) == 0, errors, warnings)


# ── Gate 5: ASSEMBLY ────────────────────────────────────────
# 流程图节点: AE (写配文+insert_config) → AF (填模板)
#            → AG (删占位/示例/格式说明) → AH (学生口吻)

def gate_assembly(workflow: dict, checklist: dict) -> tuple:
    errors, warnings = [], []
    output_dir = Path(workflow.get("output_dir", "."))
    output_docx = Path(workflow.get("output_docx", ""))

    # AE: copywriting.md 不能还是默认 stub
    copywriting_path = output_dir / "copywriting.md"
    if copywriting_path.exists():
        content = copywriting_path.read_text(encoding="utf-8")
        if "先完成 requirement_analysis.json" in content and "不要直接 run" in content:
            errors.append(
                "[Gate 5/AE] copywriting.md 仍为初始化默认内容，Agent 必须写实际报告正文"
            )
        if content.strip().startswith("# Copywriting") and len(content.strip().splitlines()) <= 10:
            warnings.append("[Gate 5/AE] copywriting.md 内容似乎过短，请确认已完成正文")

    # AE: insert_config.json 不能是空壳
    ic = _load_if_exists(output_dir / "insert_config.json") or {}
    ic_images = ic.get("images", {})
    ic_placements = ic.get("placements", {})
    if not ic_images:
        warnings.append("[Gate 5/AE] insert_config.json.images 为空，图片将无法插入")
    if not ic_placements:
        warnings.append("[Gate 5/AE] insert_config.json.placements 为空")

    # AF: 检查 docx 脚本是否仍是 stub
    for name in ("fill", "insert", "verify"):
        script_path_str = workflow.get("docx_scripts", {}).get(name, "")
        if script_path_str:
            sp = Path(script_path_str)
            if sp.exists() and "AUTO_LAB_TEMPLATE_SCRIPT_STUB = True" in sp.read_text(encoding="utf-8"):
                errors.append(
                    f"[Gate 5/AF] task_scripts/{name}_template.py 仍为 stub，"
                    f"Agent 必须根据模板结构定制化填写"
                )

    # AF: 输出 DOCX 是否存在（如果已执行 run）
    if output_docx != Path("") and not output_docx.exists():
        warnings.append(
            f"[Gate 5/AF] 输出文档尚未生成: {output_docx}，"
            f"请确认已运行 run 命令"
        )

    # AG + AH: 输出 DOCX 的文本 lint（检查残留占位、localhost、AI味等）
    if output_docx != Path("") and output_docx.exists():
        try:
            from docx import Document
            doc = Document(str(output_docx))
            paragraphs = [p.text for p in doc.paragraphs]
            # 也扫描表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.append(cell.text)
            full_text = "\n".join(paragraphs)

            # AG: 占位/示例文本
            matches = FORBIDDEN_IN_OUTPUT_RE.findall(full_text)
            if matches:
                unique_matches = list(set(m.lower() for m in matches))
                errors.append(
                    f"[Gate 5/AG] 输出文档中发现 {len(unique_matches)} 种禁止文本: "
                    + ", ".join(unique_matches[:8])
                )

            # AG: 检查 TOC 是否残留为空白或占位
            toc_suspect = any(
                phrase in full_text.lower()
                for phrase in ["目录", "table of contents", "toc"]
            )
            if toc_suspect:
                # TOC 区域后面应该有实际内容，如果紧接着是空白页则警告
                warnings.append(
                    "[Gate 5/AG] 检测到 TOC/目录标记，请确认目录是否已正确更新而非残留在空白状态"
                )

            # AH: 学生口吻基础检查（Agent 口吻常见词）
            agent_tells = [
                r"\b(?:agent|assistant|claude|gpt|llm|ai assistant|人工智能助手)\b",
                r"(?:I (?:have|will|can|decided|chose))",
                r"(?:根据.*?要求.*?(?:生成|创建|制作))",
                r"(?:本.*?(?:agent|助手|模型|系统).*?(?:生成|创建|输出))",
            ]
            import re as re_module
            for pattern in agent_tells:
                found = re_module.findall(pattern, full_text, re_module.IGNORECASE)
                if found:
                    errors.append(
                        f"[Gate 5/AH] 输出文档疑似使用 Agent 口吻，"
                        f"发现 {len(found)} 处匹配模式: {pattern[:60]}..."
                    )
                    break  # 报一次就够了

        except Exception as e:
            warnings.append(f"[Gate 5] 无法解析输出 DOCX 进行文本检查: {e}")

    return (len(errors) == 0, errors, warnings)


# ── Gate 6: DELIVERY ────────────────────────────────────────
# 流程图节点: AI (写 submission_package) → AJ (交付审查) → AK (完整性门) → AM (打包)

def gate_delivery(workflow: dict, checklist: dict) -> tuple:
    errors, warnings = [], []
    output_dir = Path(workflow.get("output_dir", "."))

    if not bool(checklist.get("submission_package_required", False)):
        return (True, [], [])

    # AI: submission_package.json
    sp = _load_if_exists(output_dir / "submission_package.json") or {}
    if not sp.get("enabled", False):
        errors.append("[Gate 6/AI] submission_package_required=true 但 submission_package.json.enabled 仍为 false")
    if not sp.get("include_paths"):
        errors.append("[Gate 6/AI] submission_package.json.include_paths 为空")

    # Naming check: must be submit.zip
    output_zip = sp.get("output_zip", "")
    zip_name = Path(output_zip).name if output_zip else ""
    if zip_name and zip_name.lower() != "submit.zip" and not sp.get("allow_custom_zip_name", False):
        errors.append(
            f"[Gate 6/AI] submission archive name must be 'submit.zip', got '{zip_name}'. "
            f"Remove extra suffixes or set allow_custom_zip_name=true."
        )

    # AK→AM: submit.zip 存在性
    submit_zip = output_dir / "submit.zip"
    if not submit_zip.exists():
        errors.append(
            f"[Gate 6/AM] submit.zip 不存在: {submit_zip}，"
            f"请运行 'python scripts/run_workflow.py package --workflow <workflow.json>'"
        )

    # submit/ folder 存在性
    submit_dir = output_dir / "submit"
    if not submit_dir.exists():
        warnings.append(
            f"[Gate 6/AM] submit/ 文件夹不存在: {submit_dir}，"
            f"请运行 'python scripts/run_workflow.py package --workflow <workflow.json>'"
        )

    # 检查 include_paths 中的文件是否真实存在
    missing_includes = []
    for item in sp.get("include_paths", []):
        p = item.get("path", "") if isinstance(item, dict) else str(item)
        if p and not Path(p).exists():
            missing_includes.append(p)
    if missing_includes:
        errors.append(
            f"[Gate 6/AJ] submission_package.json 中 {len(missing_includes)} 个文件不存在: "
            + ", ".join(missing_includes[:5])
        )

    return (len(errors) == 0, errors, warnings)


# ── Gate 调度器 ─────────────────────────────────────────────

def _load_if_exists(path: Path):
    if path.exists():
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            return None
    return None


_GATES = {
    "1": ("INIT — 环境与初始化", gate_init),
    "2": ("PRE_TASK — 前置任务", gate_pre_task),
    "3": ("EVIDENCE — 证据规划", gate_evidence),
    "4": ("GENERATION — 产物生成与审查", gate_generation),
    "5": ("ASSEMBLY — 装配与清理", gate_assembly),
    "6": ("DELIVERY — 交付打包", gate_delivery),
}


def generate_delivery_review(workflow: dict, checklist: dict) -> dict:
    """A10: Generate machine-checkable delivery_review.json."""
    output_dir = Path(workflow.get("output_dir", "."))
    output_docx = Path(workflow.get("output_docx", ""))

    review = {
        "status": "pending",
        "checks": []
    }

    # Check 1: Output DOCX exists
    docx_exists = output_docx.exists() if output_docx != Path("") else False
    review["checks"].append({
        "name": "output_docx_exists",
        "passed": docx_exists,
        "detail": str(output_docx) if docx_exists else "Output DOCX not found"
    })

    # Check 2: Images generated
    images_dir = Path(workflow.get("images_dir", str(output_dir / "generated_images")))
    planned = checklist.get("planned_figures", [])
    ai_planned = [f for f in planned if f.get("source_mode") == "ai_simulated"]
    browser_planned = [f for f in planned if f.get("source_mode") == "browser_capture"]
    diagram_planned = [f for f in planned if f.get("source_mode") == "diagram_assets"]

    for fig in ai_planned + browser_planned + diagram_planned:
        name = fig.get("name", "?")
        png_path = images_dir / f"{name}.png"
        review["checks"].append({
            "name": f"image_{name}",
            "passed": png_path.exists(),
            "detail": str(png_path) if png_path.exists() else f"Image not found: {name}"
        })

    # Check 3: Template scripts customized
    docx_scripts = workflow.get("docx_scripts", {})
    for script_name in ("fill", "insert", "verify"):
        script_path = Path(docx_scripts.get(script_name, ""))
        is_stub = False
        if script_path.exists():
            content = script_path.read_text(encoding="utf-8")
            is_stub = "AUTO_LAB_TEMPLATE_SCRIPT_STUB = True" in content
        review["checks"].append({
            "name": f"script_{script_name}_customized",
            "passed": not is_stub,
            "detail": "Stub detected" if is_stub else "Customized"
        })

    # Check 4: Submission package
    if checklist.get("submission_package_required", False):
        submit_folder = output_dir / "submit"
        submit_zip = output_dir / "submit.zip"
        review["checks"].append({
            "name": "submit_folder_exists",
            "passed": submit_folder.exists(),
            "detail": str(submit_folder) if submit_folder.exists() else "submit/ folder not found"
        })
        review["checks"].append({
            "name": "submit_zip_exists",
            "passed": submit_zip.exists(),
            "detail": str(submit_zip) if submit_zip.exists() else "submit.zip not found"
        })

    # Overall status
    all_passed = all(c["passed"] for c in review["checks"])
    review["status"] = "pass" if all_passed else "fail"
    review["total_checks"] = len(review["checks"])
    review["passed_checks"] = sum(1 for c in review["checks"] if c["passed"])

    # Write to file
    review_path = output_dir / "delivery_review.json"
    review_path.write_text(json.dumps(review, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Delivery review written: {review_path}")

    return review


def run_gates(workflow: dict, phase: str = "all") -> int:
    """运行 Phase Gate 检查。返回 0 = 全部通过，非 0 = 有错误。"""
    output_dir = Path(workflow.get("output_dir", "."))
    checklist = _load_if_exists(output_dir / "requirement_checklist.json") or {}
    requirement_analysis = _load_if_exists(output_dir / "requirement_analysis.json") or {}
    pre_task_plan = _load_if_exists(output_dir / "pre_task_plan.json") or {
        "enabled": False, "completed": False, "task_type": "",
        "objective": "", "output_summary": "", "output_paths": [], "output_artifacts": [],
    }

    gate_context = {
        "workflow": workflow,
        "output_dir": output_dir,
        "checklist": checklist,
        "requirement_analysis": requirement_analysis,
        "pre_task_plan": pre_task_plan,
    }

    phases_to_run = list(_GATES.keys()) if phase == "all" else [phase]
    total_errors = 0
    total_warnings = 0
    any_fatal = False

    print("=" * 60)
    print("auto-lab Phase Gate 检查")
    print(f"输出目录: {output_dir}")
    print(f"运行阶段: {', '.join(phases_to_run)}")
    print("=" * 60)

    for ph in phases_to_run:
        if ph not in _GATES:
            print(f"\n[Phase {ph}] 未知阶段，跳过")
            continue

        title, gate_func = _GATES[ph]
        print(f"\n{'─' * 60}")
        print(f"🚪 Gate {ph}: {title}")
        print(f"{'─' * 60}")

        try:
            if ph == "1":
                passed, errs, warns = gate_func(workflow, output_dir, checklist)
            elif ph == "2":
                passed, errs, warns = gate_func(workflow, checklist, pre_task_plan)
            elif ph == "3":
                passed, errs, warns = gate_func(workflow, checklist, requirement_analysis)
            elif ph in ("4", "5", "6"):
                passed, errs, warns = gate_func(workflow, checklist)
            else:
                passed, errs, warns = True, [], []
        except Exception as e:
            passed, errs, warns = False, [f"Gate {ph} 执行异常: {e}"], []

        for w in warns:
            print(f"  ⚠️  {w}")
            total_warnings += 1

        for e in errs:
            print(f"  ❌ {e}")
            total_errors += 1

        if passed:
            print(f"  ✅ Gate {ph} 通过")
        else:
            print(f"  🚫 Gate {ph} 未通过 — 后续阶段将跳过")
            any_fatal = True
            if phase == "all":
                print(f"  (跳过 Gate {int(ph)+1}~6，请修复后重新运行)")
                break  # 全量模式遇错即停

    print(f"\n{'=' * 60}")
    print(f"汇总: {total_errors} 错误, {total_warnings} 警告")
    if any_fatal:
        print("状态: 🚫 未通过 — 请修复上方错误后重新运行 gate")
        print("=" * 60)
        return 1
    else:
        if total_warnings > 0:
            print("状态: ✅ 通过 (有警告)")
        else:
            print("状态: ✅ 全部通过")
        print("=" * 60)
        return 0


def main():
    args = parse_args()
    workflow = load_workflow(args.workflow)
    checklist, requirement_analysis, prompt_config, browser_capture_plan, diagram_plan, video_plan, reference_template_cleanup, submission_package, insert_config, notes = validate_workflow(workflow)
    for note in notes:
        print(note)
    if args.command == "validate":
        print("Validation passed.")
        return

    if args.command == "gate":
        phase = getattr(args, "phase", "all")
        exit_code = run_gates(workflow, phase)
        sys.exit(exit_code)

    update_insert_config_from_prompts(workflow, prompt_config, insert_config)
    update_insert_config_from_browser_plan(workflow, browser_capture_plan, insert_config)
    update_insert_config_from_diagram_plan(workflow, diagram_plan, insert_config)

    if args.command == "images":
        if checklist.get("ai_images_required", False):
            run_generate_py(workflow)
        if checklist.get("browser_capture_required", False):
            run_browser_capture(workflow)
        if checklist.get("diagram_assets_required", False):
            run_diagram_assets(workflow)
        print("Image generation finished.")
        return

    if args.command == "video":
        if checklist.get("video_required", False):
            run_video_processing(workflow, video_plan)
        print("Video processing finished.")
        return

    if args.command == "package":
        if checklist.get("submission_package_required", False):
            run_submission_package(workflow)
        print("Submission packaging finished.")
        return

    if args.command == "run":
        ensure_docx_scripts_customized(workflow)
        if checklist.get("reference_template_cleanup_required", False):
            run_reference_template_cleanup(workflow, reference_template_cleanup)
        if checklist.get("ai_images_required", False):
            run_generate_py(workflow)
        if checklist.get("browser_capture_required", False):
            run_browser_capture(workflow)
        if checklist.get("diagram_assets_required", False):
            run_diagram_assets(workflow)
        if checklist.get("video_required", False):
            run_video_processing(workflow, video_plan)
        run_manifest_commands(workflow)
        if checklist.get("submission_package_required", False):
            run_submission_package(workflow)
        # A10: Generate machine-checkable delivery review
        generate_delivery_review(workflow, checklist)
        print("Workflow run finished.")


if __name__ == "__main__":
    main()
