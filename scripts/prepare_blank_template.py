import argparse
import copy
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


BODY_TEXT_STYLES = {
    "body text",
    "正文",
    "list paragraph",
    "normal",
    "caption",
    "图题",
    "表题",
}

LIST_MARK_RE = re.compile(r"^\s*(\(\d+\)|（\d+）|[0-9]+[.)、]|[a-zA-Z][.)])")
CHAPTER_RE = re.compile(r"^\s*(第\s*[一二三四五六七八九十百千万0-9]+\s*[章节章]|chapter\s+\d+)", re.I)


def parse_args():
    parser = argparse.ArgumentParser(description="Create a blank DOCX template from a filled reference document.")
    parser.add_argument("--input", required=True, help="Filled reference .docx")
    parser.add_argument("--output", required=True, help="Blank template .docx output")
    parser.add_argument("--report", help="Optional JSON cleanup report")
    parser.add_argument("--keep-levels", default="1,2", help="Heading levels to keep after first chapter, default 1,2")
    return parser.parse_args()


def style_name(paragraph):
    return (paragraph.style.name if paragraph.style is not None else "").strip()


def outline_level(paragraph):
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return None
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        return None
    value = outline.get(qn("w:val"))
    return int(value) if value is not None and value.isdigit() else None


def heading_level(paragraph):
    name = style_name(paragraph).lower().replace(" ", "")
    match = re.search(r"heading([1-9])", name)
    if match:
        return int(match.group(1))
    cn_match = re.search(r"标题([1-9一二三四五六七八九])", style_name(paragraph))
    if cn_match:
        raw = cn_match.group(1)
        cn = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}
        return cn.get(raw, int(raw) if raw.isdigit() else None)
    level = outline_level(paragraph)
    if level is not None:
        return level + 1
    text = paragraph.text.strip()
    if CHAPTER_RE.match(text):
        return 1
    return None


def has_drawing(paragraph):
    return bool(paragraph._p.xpath(".//w:drawing | .//w:pict"))


def remove_element(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def clone_paragraph_without_media(paragraph):
    cloned = copy.deepcopy(paragraph._p)
    for node in cloned.xpath(".//w:drawing | .//w:pict"):
        remove_element(node)
    return cloned


def table_starts_after_body(table, body_start_element):
    body = table._tbl.getparent()
    children = list(body)
    try:
        return children.index(table._tbl) > children.index(body_start_element)
    except ValueError:
        return False


def should_remove_paragraph(paragraph, keep_levels):
    level = heading_level(paragraph)
    if level in keep_levels:
        return False
    name = style_name(paragraph).strip().lower()
    text = paragraph.text.strip()
    if has_drawing(paragraph):
        return True
    if name in BODY_TEXT_STYLES:
        return True
    if "caption" in name or "题注" in name:
        return True
    if LIST_MARK_RE.match(text):
        return True
    if level is not None:
        return True
    if text:
        return True
    return True


def find_body_start(doc):
    for paragraph in doc.paragraphs:
        level = heading_level(paragraph)
        if level == 1:
            return paragraph._p
    return None


def cleanup(input_path: Path, output_path: Path, keep_levels):
    doc = Document(str(input_path))
    body_start = find_body_start(doc)
    if body_start is None:
        raise SystemExit("Could not find the first level-1 heading. Stop rather than guessing the body boundary.")

    body = doc._body._element
    children = list(body)
    body_start_index = children.index(body_start)
    removed = {"paragraphs": 0, "tables": 0, "media_nodes": 0, "kept_headings": 0}

    for table in list(doc.tables):
        if table_starts_after_body(table, body_start):
            remove_element(table._tbl)
            removed["tables"] += 1

    for paragraph in list(doc.paragraphs):
        element = paragraph._p
        current_children = list(body)
        if element not in current_children:
            continue
        if current_children.index(element) < body_start_index:
            continue
        level = heading_level(paragraph)
        if level in keep_levels:
            media_nodes = paragraph._p.xpath(".//w:drawing | .//w:pict")
            for node in list(media_nodes):
                remove_element(node)
                removed["media_nodes"] += 1
            removed["kept_headings"] += 1
            continue
        if should_remove_paragraph(paragraph, keep_levels):
            removed["paragraphs"] += 1
            remove_element(element)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return {
        "input": str(input_path),
        "output": str(output_path),
        "body_start_index": body_start_index,
        "keep_heading_levels": sorted(keep_levels),
        "removed": removed,
    }


def main():
    args = parse_args()
    keep_levels = {int(item.strip()) for item in args.keep_levels.split(",") if item.strip()}
    report = cleanup(Path(args.input).expanduser().resolve(), Path(args.output).expanduser().resolve(), keep_levels)
    if args.report:
        report_path = Path(args.report).expanduser().resolve()
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
