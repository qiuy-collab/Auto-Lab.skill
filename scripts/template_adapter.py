"""
Generic DOCX template adapter for auto-lab.

Provides reusable functions for common DOCX operations:
- Finding and replacing placeholder text
- Inserting images at placeholder positions
- Removing template instruction text
- Validating filled template output

This reduces the need to write task-specific scripts from scratch.
Use this as a library in task_scripts/fill_template.py, insert_images.py, etc.
"""
import json
import re
import sys
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt


PLACEHOLDER_RE = re.compile(r"\{\{(img_\d{2})\}\}")


def load_insert_config(config_path: str) -> dict:
    return json.loads(Path(config_path).read_text(encoding="utf-8"))


def find_placeholders(doc: Document) -> list[str]:
    """Find all {{img_XX}} placeholders in document paragraphs."""
    found = []
    for para in doc.paragraphs:
        matches = PLACEHOLDER_RE.findall(para.text)
        found.extend(matches)
    return sorted(set(found))


def find_paragraph_with_placeholder(doc: Document, key: str) -> Optional[int]:
    """Return the paragraph index that contains the placeholder, or None."""
    marker = "{{" + key + "}}"
    for idx, para in enumerate(doc.paragraphs):
        if marker in para.text:
            return idx
    return None


def replace_placeholder_text(doc: Document, key: str, replacement: str):
    """Replace {{key}} placeholder text in all paragraphs."""
    marker = "{{" + key + "}}"
    for para in doc.paragraphs:
        if marker in para.text:
            para.text = para.text.replace(marker, replacement)


def insert_image_after_paragraph(
    doc: Document,
    para_index: int,
    image_path: str,
    width_inches: float = 6.0,
):
    """Insert an image after the specified paragraph."""
    para = doc.paragraphs[para_index]
    run = para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))


def remove_template_instructions(doc: Document, patterns: list[str] = None):
    """Remove paragraphs matching template instruction patterns."""
    if patterns is None:
        patterns = [
            r"请在此处填写",
            r"请填写",
            r"\[TODO\]",
            r"TODO[：:]",
            r"占位",
            r"示例内容",
            r"样例",
            r"此处替换",
            r"根据模板",
        ]
    compiled = [re.compile(p, re.IGNORECASE) for p in patterns]
    to_remove = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        for pattern in compiled:
            if pattern.search(text):
                to_remove.append(para)
                break
    for para in to_remove:
        p = para._element
        p.getparent().remove(p)


def validate_output(docx_path: str, insert_config_path: str) -> dict:
    """Validate that the filled template has all expected images and no leftover placeholders."""
    doc = Document(docx_path)
    config = load_insert_config(insert_config_path)
    expected_images = set(config.get("images", {}).keys())
    found_placeholders = set(find_placeholders(doc))

    missing = expected_images - found_placeholders
    unexpected = found_placeholders - expected_images

    return {
        "valid": len(missing) == 0 and len(unexpected) == 0,
        "expected_images": sorted(expected_images),
        "missing_placeholders": sorted(missing),
        "unexpected_placeholders": sorted(unexpected),
    }


def write_fill_script(
    template_path: str,
    copywriting_path: str,
    output_path: str,
    insert_config_path: str,
    width_inches: float = 6.0,
):
    """
    Generic fill + insert script.

    1. Loads the template
    2. Replaces {{img_XX}} placeholders with image references
    3. Inserts images at placeholder positions
    4. Removes template instruction text
    5. Saves the output
    """
    doc = Document(template_path)
    insert_config = load_insert_config(insert_config_path)
    images = insert_config.get("images", {})

    for key, image_path in images.items():
        marker = "{{" + key + "}}"
        for para in doc.paragraphs:
            if marker in para.text:
                para.text = para.text.replace(marker, "")
                run = para.add_run()
                if Path(image_path).exists():
                    run.add_picture(image_path, width=Inches(width_inches))

    remove_template_instructions(doc)
    doc.save(output_path)
    print(f"Filled template saved: {output_path}")


if __name__ == "__main__":
    print("template_adapter.py is a library module.")
    print("Import it in task_scripts/ or use its functions directly.")
    print()
    print("Available functions:")
    print("  load_insert_config(path)")
    print("  find_placeholders(doc)")
    print("  replace_placeholder_text(doc, key, text)")
    print("  insert_image_after_paragraph(doc, para_index, image_path, width)")
    print("  remove_template_instructions(doc, patterns)")
    print("  validate_output(docx_path, config_path)")
    print("  write_fill_script(template, copywriting, output, config, width)")
